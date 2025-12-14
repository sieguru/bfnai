"""
Document parser service - Parse .docx files and extract structure
"""
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from typing import List, Dict, Any, Optional
from collections import Counter


def parse_document_structured(file_path: str) -> Dict[str, Any]:
    """
    Parse a .docx file and extract structured content with styles.
    Returns paragraphs with their style information.
    """
    try:
        doc = Document(file_path)
    except PackageNotFoundError:
        raise ValueError(f"Could not open document: {file_path}")

    paragraphs = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else 'Normal'

        paragraphs.append({
            'index': len(paragraphs),
            'text': text,
            'style': style_name,
        })

    return {
        'paragraphs': paragraphs,
        'metadata': {
            'paragraph_count': len(paragraphs),
        }
    }


def analyze_document_styles(file_path: str) -> List[Dict[str, Any]]:
    """
    Analyze a document and return detected styles with sample text and counts.
    """
    try:
        doc = Document(file_path)
    except PackageNotFoundError:
        raise ValueError(f"Could not open document: {file_path}")

    style_info = {}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else 'Normal'

        if style_name not in style_info:
            style_info[style_name] = {
                'style_name': style_name,
                'sample_text': text[:200],
                'occurrence_count': 0,
                'heading_level': _detect_heading_level(style_name)
            }

        style_info[style_name]['occurrence_count'] += 1

    # Sort by occurrence count
    styles = sorted(
        style_info.values(),
        key=lambda x: x['occurrence_count'],
        reverse=True
    )

    return styles


def _detect_heading_level(style_name: str) -> Optional[int]:
    """
    Try to auto-detect heading level from style name.
    Returns None if not a heading style.
    """
    style_lower = style_name.lower()

    # Common heading patterns
    heading_patterns = {
        'heading 1': 1, 'heading1': 1, 'rubrik 1': 1, 'titel': 1, 'title': 1,
        'heading 2': 2, 'heading2': 2, 'rubrik 2': 2, 'subtitle': 2,
        'heading 3': 3, 'heading3': 3, 'rubrik 3': 3,
        'heading 4': 4, 'heading4': 4, 'rubrik 4': 4,
        'heading 5': 5, 'heading5': 5, 'rubrik 5': 5,
        'heading 6': 6, 'heading6': 6, 'rubrik 6': 6,
    }

    for pattern, level in heading_patterns.items():
        if pattern in style_lower:
            return level

    return None
